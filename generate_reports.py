import os
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    subprocess.check_call(["pip", "install", "python-docx"])
    import docx

from docx import Document

def create_report(filename, title, intro, details, specific_title, specific_content, challenges):
    doc = Document()
    doc.add_heading(title, 0)
    
    doc.add_paragraph('ITD103 – MODERN QUERY LANGUAGES')
    doc.add_paragraph('Instructor: Ramat B. Hailaya')
    doc.add_paragraph('Sualden S. Sala')
    doc.add_paragraph('June 16, 2026')
    doc.add_paragraph('')
    
    doc.add_heading('1. Introduction & Objectives', level=1)
    doc.add_paragraph(intro)
    
    doc.add_heading('2. Implementation Details', level=1)
    doc.add_paragraph(details)
    
    doc.add_heading(f'3. {specific_title}', level=1)
    doc.add_paragraph(specific_content)
    
    doc.add_heading('4. Challenges Encountered & Solutions', level=1)
    doc.add_paragraph(challenges)
    
    doc.add_paragraph('(Note to Instructor: All visual outputs of the queries and server tests can be found inside the "screenshot" folder).')
    
    doc.save(filename)

# Lab 10
create_report(
    '/home/dens/Itd103/itd103-labs/lab10-12-graphql/lab10-graphql-fundamentals/Lab_10_Report.docx',
    'Lab 10: GraphQL Fundamentals',
    'The goal of this laboratory was to transition from RESTful APIs to GraphQL by building a fundamental GraphQL server. This exercise demonstrated how clients can request exact data structures, eliminating the common issues of over-fetching and under-fetching associated with traditional endpoints.',
    'A Node.js server was constructed utilizing Apollo Server. A schema was defined containing types for Book and Query. The dataset was stored in-memory. The primary books query resolver was implemented to map incoming client requests directly to the mock dataset, allowing clients to dynamically specify which fields (like title or author) they wished to retrieve.',
    'Query Execution & Testing',
    'Testing was conducted using Apollo Sandbox and cURL scripts. Queries were successfully executed to retrieve specific data fields. The environment setup script ensured the correct Node.js and npm versions were active before initializing the server on port 4000.',
    'During the testing phase, initial connection errors occurred because the test scripts attempted to query the server before it had fully started. This was resolved by implementing a brief asynchronous delay in the test scripts to ensure the Apollo Server was fully bound to port 4000 before dispatching the HTTP POST request.'
)

# Lab 11
create_report(
    '/home/dens/Itd103/itd103-labs/lab10-12-graphql/lab11-advanced-graphql/Lab_11_Report.docx',
    'Lab 11: Advanced GraphQL',
    'The objective of this laboratory was to implement advanced GraphQL concepts, specifically focusing on nested resolvers and API authentication. This demonstrated how GraphQL handles complex, deeply related data hierarchies and how security can be enforced at the endpoint level.',
    'The schema from the previous lab was expanded to include an Author type with a one-to-many relationship to the Book type. Nested resolvers were engineered so that querying an author automatically resolved and fetched their associated books array. Additionally, a secretData query was protected by an authentication check utilizing HTTP headers.',
    'Authentication & Authorization',
    'The Apollo Server context function was utilized to extract the Authorization header from incoming requests. If a valid bearer token (my-secret-token) was not provided, the server actively threw an AuthenticationError, successfully blocking unauthorized access to the secretData resolver.',
    'When designing the nested resolvers, resolving the bidirectional relationship between Authors and Books initially caused confusion regarding the parent object parameter. This was solved by carefully mapping the parent.id of the Author to the authorId property in the books dataset within the Author.books resolver function.'
)

# Lab 12
create_report(
    '/home/dens/Itd103/itd103-labs/lab10-12-graphql/lab12-graphql-multiple-sources/Lab_12_Report.docx',
    'Lab 12: Microservices Federation',
    'The final GraphQL laboratory focused on scaling architectures via Microservices Federation. The objective was to combine multiple independent GraphQL APIs (subgraphs) into a single, unified supergraph using an API Gateway, allowing seamless querying across distributed services.',
    'Three distinct Node.js services were created using Docker containers: a users-service, a books-service, and a gateway. Apollo Federation v2 directives (@key) were utilized in the subgraphs to define shared entities. The API Gateway was configured using IntrospectAndCompose to automatically discover the subgraphs and route incoming client queries to the appropriate underlying microservice.',
    'Unified Supergraph Queries',
    'A unified query was executed against the gateway on port 4000, requesting a user\'s name alongside the specific titles of their borrowed books. The gateway successfully decomposed the query, fetched the user data from the users service, resolved the book references from the books service, and stitched the JSON response together.',
    'During deployment, Fedora\'s SELinux security policies blocked the Docker containers from accessing the host machine\'s source code, causing module loading failures. This was resolved by appending the :Z flag to the Docker Compose volume mounts, instructing SELinux to correctly label the directories for container access. Furthermore, a schema composition error involving the @external directive in Federation v2 was fixed by correctly utilizing the @key directive on both subgraphs.'
)

print("Reports generated successfully!")
